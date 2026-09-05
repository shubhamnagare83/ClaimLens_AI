"""
ClaimLens AI — Advanced Document Service
Multi-engine PDF reader (PyMuPDF + PyPDF2 fallback), DOCX, TXT parser,
AI document classification, and real-world insurance entity extraction.
"""
import io
import re
from pathlib import Path
from typing import Dict, List, Any, Tuple, Optional
from backend.services import gemini_service

# Supported extensions
SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.txt', '.csv', '.json'}


def read_document(filepath: str) -> Tuple[str, str]:
    """Read a document from disk and return (content, doc_type)."""
    path = Path(filepath)
    ext = path.suffix.lower()

    if ext == '.txt':
        return _read_txt(path), 'txt'
    elif ext == '.pdf':
        return _read_pdf(path), 'pdf'
    elif ext == '.docx':
        return _read_docx(path), 'docx'
    else:
        try:
            return path.read_text(encoding='utf-8'), ext.lstrip('.')
        except:
            return "", "unknown"


def read_content(content: bytes, filename: str) -> str:
    """Read plain text content from bytes based on filename extension."""
    parsed = parse_document_bytes(content, filename)
    return parsed.get("text", "")


def parse_document_bytes(content: bytes, filename: str) -> Dict[str, Any]:
    """
    Parse uploaded document bytes into clean text and rich structural metadata.
    """
    ext = Path(filename).suffix.lower()
    text = ""
    pages = 1
    metadata = {}
    engine = "generic"

    if ext == '.pdf':
        text, pages, metadata, engine = _read_pdf_with_metadata(content)
    elif ext == '.docx':
        text = _read_docx_bytes(content)
        engine = "python-docx"
    elif ext == '.txt' or ext == '.csv' or ext == '.json':
        try:
            text = content.decode('utf-8')
        except UnicodeDecodeError:
            text = content.decode('latin-1', errors='replace')
        engine = "text-decoder"
    else:
        try:
            text = content.decode('utf-8')
            engine = "fallback-utf8"
        except Exception:
            text = content.decode('latin-1', errors='replace')
            engine = "fallback-latin1"

    # Auto-classify document type
    classification = classify_document_type(text, filename)

    # Extract structured entities
    entities = extract_entities_from_text(text, filename)

    return {
        "filename": filename,
        "extension": ext,
        "file_size": len(content),
        "page_count": pages,
        "engine": engine,
        "text": text,
        "document_type": classification["detected_type"],
        "classification_confidence": classification["confidence"],
        "classification_reasons": classification["reasons"],
        "entities": entities,
        "metadata": metadata
    }


def _read_pdf_with_metadata(content: bytes) -> Tuple[str, int, Dict[str, Any], str]:
    """
    Primary: PyMuPDF (fitz) for accurate text layout and table parsing.
    Secondary: PyPDF2 as fallback.
    """
    # 1. Try PyMuPDF
    try:
        import fitz
        doc = fitz.open(stream=content, filetype="pdf")
        page_texts = []
        for i in range(len(doc)):
            page = doc[i]
            page_texts.append(page.get_text("text"))
        full_text = "\n\n".join(page_texts)
        meta = doc.metadata or {}
        pages = len(doc)
        doc.close()
        return full_text.strip(), pages, meta, "PyMuPDF (fitz)"
    except Exception as e_fitz:
        # 2. Try PyPDF2
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(io.BytesIO(content))
            page_texts = []
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    page_texts.append(t)
            full_text = "\n\n".join(page_texts)
            pages = len(reader.pages)
            return full_text.strip(), pages, {}, "PyPDF2"
        except Exception as e_pypdf:
            return f"[PDF Parsing Error: fitz={e_fitz}, pypdf={e_pypdf}]", 0, {}, "error"


def _read_pdf(path: Path) -> str:
    try:
        content = path.read_bytes()
        text, _, _, _ = _read_pdf_with_metadata(content)
        return text
    except Exception as e:
        return f"[PDF parsing error: {e}]"


def _read_docx(path: Path) -> str:
    try:
        from docx import Document
        doc = Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n".join(paragraphs)
    except Exception as e:
        return f"[DOCX parsing error: {e}]"


def _read_docx_bytes(content: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n".join(paragraphs)
    except Exception as e:
        return f"[DOCX parsing error: {e}]"


def _read_txt(path: Path) -> str:
    try:
        return path.read_text(encoding='utf-8')
    except:
        return path.read_text(encoding='latin-1', errors='replace')


def classify_document_type(text: str, filename: str = "") -> Dict[str, Any]:
    """
    Intelligently classify the document type based on contents and filename cues.
    Returns:
      detected_type: 'claim_form' | 'repair_estimate' | 'fir' | 'surveyor_report' | 'incident_statement'
      confidence: float (0.0 to 1.0)
      reasons: list of matching criteria
    """
    t = text.lower()
    fn = filename.lower()
    reasons = []
    scores = {
        "claim_form": 0.0,
        "repair_estimate": 0.0,
        "fir": 0.0,
        "surveyor_report": 0.0,
        "incident_statement": 0.0
    }

    # 1. Police FIR cues
    if "first information report" in t or "fir no" in t or "cr.p.c" in t or "ipc" in t or "police station" in t or "complainant" in t:
        scores["fir"] += 0.65
        reasons.append("Contains Police FIR legal terminology (IPC / Cr.P.C / Station)")
    if "stolen" in t or "theft" in t:
        scores["fir"] += 0.20
    if "fir" in fn:
        scores["fir"] += 0.35

    # 2. Repair Estimate cues
    if "estimate" in t or "quotation" in t or "garage" in t or "workshop" in t or "labour" in t or "gstin" in t or "spare parts" in t:
        scores["repair_estimate"] += 0.60
        reasons.append("Contains workshop estimate terminology (Labour / Parts / GSTIN)")
    if "total estimate" in t or "subtotal" in t or "parts subtotal" in t or "denting" in t or "painting" in t:
        scores["repair_estimate"] += 0.25
        reasons.append("Contains itemized breakdown headers")
    if "estimate" in fn or "invoice" in fn or "quote" in fn or "repair" in fn:
        scores["repair_estimate"] += 0.35

    # 3. Claim Form cues
    if "claim form" in t or "claim intimation" in t or "driver details" in t or "policyholder" in t or "insured declaration" in t:
        scores["claim_form"] += 0.65
        reasons.append("Contains official motor claim form headers and insured declaration")
    if "policy number" in t or "nature of loss" in t or "driving license" in t:
        scores["claim_form"] += 0.20
    if "claim" in fn and "form" in fn:
        scores["claim_form"] += 0.35

    # 4. Surveyor Report cues
    if "surveyor" in t or "loss assessment" in t or "depreciation" in t or "salvage" in t or "net assessed" in t:
        scores["surveyor_report"] += 0.70
        reasons.append("Contains independent motor surveyor loss assessment cues")
    if "surveyor" in fn:
        scores["surveyor_report"] += 0.35

    # 5. Incident Statement cues
    if "i was driving" in t or "statement of" in t or "narrative" in t or "incident description" in t or "suddenly" in t:
        scores["incident_statement"] += 0.50
        reasons.append("First-person customer narrative and incident statement pattern")
    if "statement" in fn or "incident" in fn or "customer" in fn:
        scores["incident_statement"] += 0.30

    # Pick top category
    best_type = max(scores, key=scores.get)
    best_score = min(0.99, max(0.50, scores[best_type]))

    if scores[best_type] == 0:
        best_type = "incident_statement"
        best_score = 0.50
        reasons.append("Default fallback classification")

    return {
        "detected_type": best_type,
        "confidence": round(best_score, 2),
        "reasons": reasons
    }


def extract_entities_from_text(text: str, filename: str = "") -> Dict[str, Any]:
    """
    Extract key motor claim entities from unstructured text using regex + NLP heuristics.
    Also parses line items and validates invoice math consistency.
    """
    entities = {
        "vehicle_registration": None,
        "policy_number": None,
        "customer_name": None,
        "contact_mobile": None,
        "vehicle_type": None,
        "vehicle_make_model": None,
        "incident_date": None,
        "incident_time": None,
        "incident_location": None,
        "repair_estimate_total": None,
        "repair_estimate_formatted": None,
        "fir_number": None,
        "keys_status": None,
        "damage_parts": [],
        "line_items": [],
        "math_consistency": {"is_valid": True, "parts_sum": 0, "discrepancy": 0},
        "flags": []
    }

    # 1. Vehicle Registration (Indian standard: MH01RB1289, KA01MJ9082, DL01AZ9988, etc.)
    reg_matches = re.findall(r'\b([A-Z]{2}[ -]?[0-9]{1,2}[ -]?[A-Z]{1,3}[ -]?[0-9]{4})\b', text)
    if reg_matches:
        # Normalize (remove spaces/dashes)
        normalized = reg_matches[0].replace(' ', '').replace('-', '').upper()
        entities["vehicle_registration"] = normalized

    # 2. Policy Number (POL-2024-xxx or similar)
    pol_match = re.search(r'(?:Policy\s*(?:Number|No\.?))\s*[:\s]*([A-Z0-9\-_]{6,20})', text, re.IGNORECASE)
    if pol_match:
        entities["policy_number"] = pol_match.group(1).strip()

    # 3. Customer / Complainant Name
    name_match = re.search(r'(?:Insured\s*Name|Customer\s*Name|Complainant\s*(?:Name)?|Informant\s*Name|Policyholder)\s*[:\s]*([A-Z][a-zA-Z\s]{2,30})(?:\n|,|$)', text, re.IGNORECASE)
    if name_match:
        entities["customer_name"] = name_match.group(1).strip()

    # 4. Contact Mobile
    mobile_match = re.search(r'(?:Mobile|Contact|Phone|Ph)\s*[:\s]*([+]?[0-9\s\-]{10,15})', text, re.IGNORECASE)
    if mobile_match:
        entities["contact_mobile"] = mobile_match.group(1).strip()

    # 5. Vehicle Type & Make Model
    if re.search(r'two[-\s]?wheeler|motorcycle|scooter|bike|enfield|activa', text, re.IGNORECASE):
        entities["vehicle_type"] = "Two-Wheeler"
    elif re.search(r'car|sedan|hatchback|suv|honda city|swift|creta|ertiga', text, re.IGNORECASE):
        entities["vehicle_type"] = "Car"

    model_match = re.search(r'(?:Make\s*&?\s*Model|Vehicle\s*Model)\s*[:\s]*([A-Za-z0-9\s\.\-]{3,35})(?:\n|,|$)', text, re.IGNORECASE)
    if model_match:
        entities["vehicle_make_model"] = model_match.group(1).strip()

    # 6. Incident Date & Time
    date_match = re.search(r'(?:Date\s*of\s*Incident|Incident\s*Date|Date\s*of\s*Occurrence|Date)\s*[:\s]*([0-9]{1,2}[/\-\.][0-9]{1,2}[/\-\.][0-9]{2,4}|[0-9]{1,2}(?:st|nd|rd|th)?\s+[A-Za-z]+\s+[0-9]{4})', text, re.IGNORECASE)
    if date_match:
        entities["incident_date"] = date_match.group(1).strip()

    time_match = re.search(r'(?:Time\s*of\s*Incident|Incident\s*Time|Approx\.?\s*Time|Time)\s*[:\s]*([0-9]{1,2}:[0-9]{2}(?:\s*(?:AM|PM|HRS))?)', text, re.IGNORECASE)
    if time_match:
        entities["incident_time"] = time_match.group(1).strip()

    # 7. Incident Location
    loc_match = re.search(r'(?:Incident\s*Location|Place\s*of\s*Occurrence|Location)\s*[:\s]*([A-Za-z0-9\s,\.\-]{4,60})(?:\n|$)', text, re.IGNORECASE)
    if loc_match:
        entities["incident_location"] = loc_match.group(1).strip()

    # 8. Police FIR Number
    fir_match = re.search(r'(?:FIR\s*(?:NO|Number|#)?)\s*[:\s]*([A-Z0-9\-_/]{4,25})', text, re.IGNORECASE)
    if fir_match:
        entities["fir_number"] = fir_match.group(1).strip()

    # 9. Repair Estimate Total & Currency
    total_match = re.search(r'(?:TOTAL\s*ESTIMATE|Total\s*(?:Amount|Cost|Repair|Estimate)?)\s*[:\s]*(?:Rs\.?|INR|₹)?\s*([0-9,]+(?:\.[0-9]{2})?)', text, re.IGNORECASE)
    if total_match:
        clean_num_str = total_match.group(1).replace(',', '')
        try:
            val = float(clean_num_str)
            entities["repair_estimate_total"] = val
            entities["repair_estimate_formatted"] = f"₹{val:,.2f}"
        except ValueError:
            pass

    # 10. Damage parts parsing
    damage_keywords = [
        "front bumper", "rear bumper", "headlamp", "headlight", "radiator grille",
        "fender", "bonnet", "windshield", "door panel", "tail lamp", "chassis",
        "suspension", "steering rack", "airbag", "denting", "painting"
    ]
    for kw in damage_keywords:
        if kw in text.lower():
            entities["damage_parts"].append(kw.title())

    # 11. Parse itemized line items (e.g., from repair estimates)
    # Looking for lines like: 01 Front Bumper Assembly ... 12,500.00
    line_item_pattern = re.findall(r'(\d{1,2})\s+([A-Za-z0-9\s\(\)/\-]{5,40}?)\s+(Part|Labour|Paint|Job)\s+(?:[0-9\sA-Za-z]+)?\s+([0-9,]+(?:\.[0-9]{2})?)', text)
    calc_sum = 0.0
    for item in line_item_pattern:
        amount_val = float(item[3].replace(',', ''))
        calc_sum += amount_val
        entities["line_items"].append({
            "index": item[0],
            "description": item[1].strip(),
            "category": item[2],
            "amount": amount_val,
            "amount_formatted": f"₹{amount_val:,.2f}"
        })

    # Math consistency check if line items exist
    if entities["line_items"] and entities["repair_estimate_total"]:
        diff = abs(calc_sum - entities["repair_estimate_total"])
        entities["math_consistency"] = {
            "is_valid": diff < 1.0,
            "parts_sum": calc_sum,
            "parts_sum_formatted": f"₹{calc_sum:,.2f}",
            "discrepancy": round(diff, 2)
        }
        if diff >= 1.0:
            entities["flags"].append(f"Mathematical discrepancy: Line items sum to ₹{calc_sum:,.2f} but reported total is {entities['repair_estimate_formatted']}")

    # 12. Commercial use warning flag
    if re.search(r'commercial|taxi|yellow board|yellow plate|fleet', text, re.IGNORECASE):
        entities["flags"].append("Potential Commercial Vehicle / Taxi usage detected — check private policy exclusion!")

    return entities
